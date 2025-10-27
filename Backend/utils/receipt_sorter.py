import os
import re
from PIL import Image, ImageOps, ImageEnhance, ImageFilter
import pytesseract
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import logging
import cv2
import numpy as np

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration
CONFIG = {
    'TESSERACT_PATH': r"C:\Users\divin\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
    'INPUT_FOLDER': "receipts",
    'OUTPUT_DOC': "Receipts_Sorted.docx",
    'IMAGE_WIDTH': Inches(2.8),
    'IMAGE_HEIGHT': Inches(3.5),
    'PADDING_COLOR': (255, 255, 255),
    'TEMP_FOLDER': "temp_processed"
}

pytesseract.pytesseract.tesseract_cmd = CONFIG['TESSERACT_PATH']

def create_temp_folder():
    """Create temporary folder for processed images."""
    if not os.path.exists(CONFIG['TEMP_FOLDER']):
        os.makedirs(CONFIG['TEMP_FOLDER'])

def preprocess_for_ocr(image_path):
    """
    Advanced image preprocessing for better OCR on low-quality images.
    Returns list of preprocessed PIL Images to try.
    """
    preprocessed_images = []
    
    try:
        # Read image with OpenCV for advanced processing
        img_cv = cv2.imread(image_path)
        if img_cv is None:
            # Fallback to PIL if OpenCV fails
            return [Image.open(image_path)]
        
        # 1. Original grayscale
        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
        preprocessed_images.append(("grayscale", Image.fromarray(gray)))
        
        # 2. Adaptive thresholding (great for uneven lighting)
        adaptive_thresh = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
        )
        preprocessed_images.append(("adaptive_thresh", Image.fromarray(adaptive_thresh)))
        
        # 3. Otsu's thresholding (automatic optimal threshold)
        _, otsu = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        preprocessed_images.append(("otsu", Image.fromarray(otsu)))
        
        # 4. Denoising + sharpening
        denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
        kernel = np.array([[-1,-1,-1], [-1,9,-1], [-1,-1,-1]])
        sharpened = cv2.filter2D(denoised, -1, kernel)
        preprocessed_images.append(("denoised_sharp", Image.fromarray(sharpened)))
        
        # 5. High contrast + sharpening
        pil_img = Image.fromarray(gray)
        contrast_enhancer = ImageEnhance.Contrast(pil_img)
        high_contrast = contrast_enhancer.enhance(2.5)
        sharpness_enhancer = ImageEnhance.Sharpness(high_contrast)
        sharp = sharpness_enhancer.enhance(2.0)
        preprocessed_images.append(("high_contrast_sharp", sharp))
        
        # 6. Morphological operations (remove noise, enhance text)
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        morph = cv2.morphologyEx(gray, cv2.MORPH_CLOSE, kernel)
        _, morph_thresh = cv2.threshold(morph, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        preprocessed_images.append(("morphological", Image.fromarray(morph_thresh)))
        
        # 7. Bilateral filter (preserve edges while reducing noise)
        bilateral = cv2.bilateralFilter(gray, 9, 75, 75)
        _, bilateral_thresh = cv2.threshold(bilateral, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        preprocessed_images.append(("bilateral", Image.fromarray(bilateral_thresh)))
        
        # 8. Increased size for better OCR (upscale by 2x)
        height, width = gray.shape
        upscaled = cv2.resize(gray, (width * 2, height * 2), interpolation=cv2.INTER_CUBIC)
        _, upscaled_thresh = cv2.threshold(upscaled, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        preprocessed_images.append(("upscaled", Image.fromarray(upscaled_thresh)))
        
    except Exception as e:
        logger.warning(f"Error in advanced preprocessing: {e}, using basic preprocessing")
        # Fallback to PIL-only preprocessing
        with Image.open(image_path) as img:
            preprocessed_images.append(("original", img.copy()))
            preprocessed_images.append(("grayscale_pil", img.convert('L')))
            preprocessed_images.append(("enhanced_pil", ImageOps.autocontrast(img.convert('L'))))
    
    return preprocessed_images

def process_image(image_path, target_width=CONFIG['IMAGE_WIDTH'], target_height=CONFIG['IMAGE_HEIGHT']):
    """Process image to have consistent size with white padding."""
    try:
        with Image.open(image_path) as img:
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            target_width_px = int(target_width.inches * 150)
            target_height_px = int(target_height.inches * 150)
            
            img_ratio = img.width / img.height
            target_ratio = target_width_px / target_height_px
            
            if img_ratio > target_ratio:
                new_width = target_width_px
                new_height = int(target_width_px / img_ratio)
            else:
                new_height = target_height_px
                new_width = int(target_height_px * img_ratio)
            
            img_resized = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            new_img = Image.new('RGB', (target_width_px, target_height_px), CONFIG['PADDING_COLOR'])
            
            x = (target_width_px - new_width) // 2
            y = (target_height_px - new_height) // 2
            new_img.paste(img_resized, (x, y))
            
            filename = os.path.basename(image_path)
            processed_path = os.path.join(CONFIG['TEMP_FOLDER'], f"processed_{filename}")
            new_img.save(processed_path, quality=95, dpi=(150, 150))
            
            return processed_path
            
    except Exception as e:
        logger.error(f"Error processing image {image_path}: {e}")
        return image_path

def parse_date_strict(date_string):
    """
    Strictly parse date string and return datetime object.
    Returns None if parsing fails or date is invalid.
    """
    date_string = date_string.strip()
    
    formats = [
        # Numeric formats with time
        "%d-%m-%Y %H:%M:%S",
        "%d/%m/%Y %H:%M:%S",
        "%m-%d-%Y %H:%M:%S",
        "%m/%d/%Y %H:%M:%S",
        "%Y-%m-%d %H:%M:%S",
        "%d-%m-%Y %H:%M",
        "%d/%m/%Y %H:%M",
        "%m-%d-%Y %H:%M",
        "%m/%d/%Y %H:%M",
        
        # Numeric formats without time
        "%d-%m-%Y",
        "%d/%m/%Y",
        "%m-%d-%Y",
        "%m/%d/%Y",
        "%Y-%m-%d",
        
        # Month name formats with time
        "%b %d, %Y %H:%M:%S",
        "%B %d, %Y %H:%M:%S",
        "%d %b %Y %H:%M:%S",
        "%d %B %Y %H:%M:%S",
        "%b %dst, %Y %H:%M:%S",
        "%b %dnd, %Y %H:%M:%S",
        "%b %drd, %Y %H:%M:%S",
        "%b %dth, %Y %H:%M:%S",
        
        # Month name formats without time
        "%b %d, %Y",
        "%B %d, %Y",
        "%d %b %Y",
        "%d %B %Y",
        "%b %dst, %Y",
        "%b %dnd, %Y",
        "%b %drd, %Y",
        "%b %dth, %Y",
        
        # With day names
        "%A, %B %d, %Y",
        "%a, %b %d, %Y",
    ]
    
    for fmt in formats:
        try:
            parsed_date = datetime.strptime(date_string, fmt)
            if 2020 <= parsed_date.year <= 2026:
                return parsed_date
        except ValueError:
            continue
    
    return None

def extract_date_from_text(text, debug_filename=""):
    """
    Extract date from text using strict pattern matching.
    Returns the most reliable date found.
    """
    text = re.sub(r'\s+', ' ', text.strip())
    
    if debug_filename:
        logger.info(f"[{debug_filename}] Analyzing text: {text[:300]}...")
    
    found_dates = []
    
    # Pattern 1: Numeric dates with various separators
    numeric_pattern = r'\b(\d{1,2})[-/](\d{1,2})[-/](20\d{2})\b'
    for match in re.finditer(numeric_pattern, text):
        full_match = match.group(0)
        d1, d2, year = match.groups()
        
        for date_str in [f"{d1}-{d2}-{year}", f"{d2}-{d1}-{year}"]:
            parsed = parse_date_strict(date_str)
            if parsed:
                found_dates.append({
                    'date': parsed,
                    'match': full_match,
                    'confidence': 'medium',
                    'type': 'numeric'
                })
                break
    
    # Pattern 2: ISO format
    iso_pattern = r'\b(20\d{2})[-/](\d{1,2})[-/](\d{1,2})\b'
    for match in re.finditer(iso_pattern, text):
        full_match = match.group(0)
        parsed = parse_date_strict(full_match)
        if parsed:
            found_dates.append({
                'date': parsed,
                'match': full_match,
                'confidence': 'high',
                'type': 'iso'
            })
    
    # Pattern 3: Month name formats
    month_pattern = r'\b(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|June?|July?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+(\d{1,2})(?:st|nd|rd|th)?,?\s+(20\d{2})\b'
    for match in re.finditer(month_pattern, text, re.IGNORECASE):
        full_match = match.group(0)
        parsed = parse_date_strict(full_match)
        if parsed:
            found_dates.append({
                'date': parsed,
                'match': full_match,
                'confidence': 'high',
                'type': 'month_name'
            })
    
    # Pattern 4: Month name with time
    month_time_pattern = r'\b(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|June?|July?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+(\d{1,2})(?:st|nd|rd|th)?,?\s+(20\d{2})\s+(\d{1,2}):(\d{2})(?::(\d{2}))?\b'
    for match in re.finditer(month_time_pattern, text, re.IGNORECASE):
        full_match = match.group(0)
        parsed = parse_date_strict(full_match)
        if parsed:
            found_dates.append({
                'date': parsed,
                'match': full_match,
                'confidence': 'very_high',
                'type': 'month_name_time'
            })
    
    # Pattern 5: Day name + month name
    day_month_pattern = r'\b(?:Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday),?\s+(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|June?|July?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+(\d{1,2}),?\s+(20\d{2})\b'
    for match in re.finditer(day_month_pattern, text, re.IGNORECASE):
        full_match = match.group(0)
        parsed = parse_date_strict(full_match)
        if parsed:
            found_dates.append({
                'date': parsed,
                'match': full_match,
                'confidence': 'very_high',
                'type': 'day_month_name'
            })
    
    if debug_filename:
        logger.info(f"[{debug_filename}] Found {len(found_dates)} potential dates:")
        for fd in found_dates:
            logger.info(f"  - '{fd['match']}' -> {fd['date'].strftime('%B %d, %Y')} (confidence: {fd['confidence']}, type: {fd['type']})")
    
    if not found_dates:
        logger.warning(f"[{debug_filename}] ❌ No valid dates found")
        return "Unknown Date"
    
    # Prioritize dates by confidence level
    confidence_order = {'very_high': 0, 'high': 1, 'medium': 2, 'low': 3}
    found_dates.sort(key=lambda x: (confidence_order.get(x['confidence'], 4), text.index(x['match'])))
    
    best_date = found_dates[0]
    result = best_date['date'].strftime("%B %d, %Y")
    
    logger.info(f"[{debug_filename}] ✅ Selected date: '{best_date['match']}' -> {result}")
    
    return result

def extract_date_from_image(image_path):
    """Extract date from receipt image using OCR with advanced preprocessing."""
    filename = os.path.basename(image_path)
    
    try:
        # OCR configurations to try
        configs = [
            r'--oem 3 --psm 6',  # Uniform block of text
            r'--oem 3 --psm 4',  # Single column
            r'--oem 3 --psm 3',  # Fully automatic
            r'--oem 3 --psm 11', # Sparse text
            r'--oem 1 --psm 6',  # LSTM only
        ]
        
        all_text = ""
        
        # Get preprocessed images
        preprocessed_images = preprocess_for_ocr(image_path)
        
        logger.info(f"[{filename}] Testing {len(preprocessed_images)} preprocessing methods × {len(configs)} OCR configs = {len(preprocessed_images) * len(configs)} combinations")
        
        # Try each preprocessed image with each config
        for img_name, img_variant in preprocessed_images:
            for config_idx, config in enumerate(configs):
                try:
                    text = pytesseract.image_to_string(img_variant, config=config)
                    all_text += " " + text
                    
                    logger.info(f"[{filename}] OCR ({img_name}-config{config_idx}): {text[:100].strip()}...")
                    
                    date_result = extract_date_from_text(text, f"{filename}-{img_name}-{config_idx}")
                    if date_result != "Unknown Date":
                        logger.info(f"[{filename}] 🎯 SUCCESS with {img_name}-config{config_idx}")
                        return date_result
                        
                except Exception as e:
                    logger.warning(f"[{filename}] OCR failed for {img_name}-config{config_idx}: {e}")
                    continue
        
        # Final attempt with all combined text
        logger.info(f"[{filename}] Trying combined text analysis (last resort)...")
        final_result = extract_date_from_text(all_text, f"{filename}-combined")
        return final_result
            
    except Exception as e:
        logger.error(f"Error extracting date from {image_path}: {e}")
        return "Unknown Date"

def set_table_borders(table):
    """Remove table borders for cleaner look."""
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    
    tbl.tblPr.append(tblBorders)

def create_receipt_document(receipts_by_date):
    """Create Word document with sorted receipts."""
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    first_page = True
    
    for date_str, file_paths in sorted(receipts_by_date.items()):
        logger.info(f"Processing date group: {date_str} ({len(file_paths)} receipts)")
        
        if not first_page:
            doc.add_page_break()
        first_page = False
        
        header = doc.add_heading(date_str, level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        for i in range(0, len(file_paths), 4):
            group = file_paths[i:i+4]
            table = doc.add_table(rows=2, cols=2)
            set_table_borders(table)
            table.autofit = False
            table.allow_autofit = False
            
            for idx, img_path in enumerate(group):
                row_idx = idx // 2
                col_idx = idx % 2
                cell = table.cell(row_idx, col_idx)
                
                try:
                    processed_img_path = process_image(img_path)
                    cell.text = ''
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(processed_img_path, width=CONFIG['IMAGE_WIDTH'])
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logger.info(f"Added image: {os.path.basename(img_path)}")
                    
                except Exception as e:
                    logger.error(f"Error adding image {img_path}: {e}")
                    cell.text = f"Error loading {os.path.basename(img_path)}"
            
            if i + 4 < len(file_paths):
                doc.add_paragraph()
        
        doc.add_paragraph()
    
    return doc

def cleanup_temp_files():
    """Clean up temporary processed images."""
    try:
        if os.path.exists(CONFIG['TEMP_FOLDER']):
            for file in os.listdir(CONFIG['TEMP_FOLDER']):
                os.remove(os.path.join(CONFIG['TEMP_FOLDER'], file))
            os.rmdir(CONFIG['TEMP_FOLDER'])
            logger.info("Cleaned up temporary files")
    except Exception as e:
        logger.warning(f"Could not clean up temporary files: {e}")

def main():
    """Main function to process receipts and create Word document."""
    logger.info("Starting receipt processing...")
    
    if not os.path.exists(CONFIG['INPUT_FOLDER']):
        logger.error(f"Input folder '{CONFIG['INPUT_FOLDER']}' not found!")
        return
    
    create_temp_folder()
    
    image_files = []
    for file in os.listdir(CONFIG['INPUT_FOLDER']):
        if file.lower().endswith(('.png', '.jpg', '.jpeg')):
            image_files.append(os.path.join(CONFIG['INPUT_FOLDER'], file))
    
    if not image_files:
        logger.error("No image files found in input folder!")
        return
    
    logger.info(f"Found {len(image_files)} receipt images")
    
    receipts_by_date = {}
    
    for img_path in image_files:
        try:
            logger.info(f"\n{'='*60}")
            logger.info(f"Processing: {os.path.basename(img_path)}")
            logger.info(f"{'='*60}")
            date_str = extract_date_from_image(img_path)
            
            if date_str not in receipts_by_date:
                receipts_by_date[date_str] = []
            receipts_by_date[date_str].append(img_path)
            
            logger.info(f"✅ Assigned '{date_str}' to {os.path.basename(img_path)}")
            
        except Exception as e:
            logger.error(f"Error processing {img_path}: {e}")
            if "Unknown Date" not in receipts_by_date:
                receipts_by_date["Unknown Date"] = []
            receipts_by_date["Unknown Date"].append(img_path)
    
    logger.info("\n" + "="*60)
    logger.info("PROCESSING SUMMARY")
    logger.info("="*60)
    for date_str, files in sorted(receipts_by_date.items()):
        logger.info(f"{date_str}: {len(files)} receipts")
        for f in files:
            logger.info(f"  - {os.path.basename(f)}")
    logger.info("="*60 + "\n")
    
    logger.info("Creating Word document...")
    doc = create_receipt_document(receipts_by_date)
    doc.save(CONFIG['OUTPUT_DOC'])
    logger.info(f"✅ Word document saved as '{CONFIG['OUTPUT_DOC']}'")
    
    cleanup_temp_files()
    logger.info("Processing complete!")

if __name__ == "__main__":
    main()

    # Fast API backend code
def process_receipts(image_paths, output_doc=None):
    """
    image_paths: list of absolute file paths of images to process
    output_doc: path to save the resulting docx (optional)
    Returns path to saved docx.
    """
    # ensure temp folder exists
    create_temp_folder()

    receipts_by_date = {}
    for img_path in image_paths:
        try:
            logger.info(f"Processing: {os.path.basename(img_path)}")
            date_str = extract_date_from_image(img_path)
            if date_str not in receipts_by_date:
                receipts_by_date[date_str] = []
            receipts_by_date[date_str].append(img_path)
            logger.info(f"Assigned date '{date_str}' to {os.path.basename(img_path)}")
        except Exception as e:
            logger.error(f"Error processing {img_path}: {e}")
            receipts_by_date.setdefault("Unknown Date", []).append(img_path)

    # create document
    doc = create_receipt_document(receipts_by_date)

    # output path
    if not output_doc:
        os.makedirs(os.path.join(os.path.dirname(__file__), '..', 'output'), exist_ok=True)
        output_doc = os.path.join(os.path.dirname(__file__), '..', 'output', CONFIG.get('OUTPUT_DOC', 'Receipts_Sorted.docx'))

    doc.save(output_doc)
    logger.info(f"Saved Word doc to: {output_doc}")

    # cleanup temp files (optional - we keep it)
    cleanup_temp_files()

    return os.path.abspath(output_doc)


#new code SSE
import shutil

def process_receipts_with_sse(image_paths, output_doc, job_id, send_progress, job_tmp_dir):
    """
    Process receipts with real-time SSE progress updates.
    This wraps the existing functions with progress reporting.
    
    Args:
        image_paths: List of image file paths
        output_doc: Path to save the output document
        job_id: Unique job identifier
        send_progress: Callback function to send progress updates
        job_tmp_dir: Temporary directory to cleanup
    """
    try:
        # Send initial status
        send_progress(job_id, f"🚀 Starting receipt processing for {len(image_paths)} files...")
        
        # Ensure temp folder exists
        create_temp_folder()
        
        receipts_by_date = {}
        
        # Process each image
        for idx, img_path in enumerate(image_paths, 1):
            filename = os.path.basename(img_path)
            
            try:
                # Send processing update
                percentage = int((idx / len(image_paths)) * 100)
                send_progress(job_id, f"⏳ [{percentage}%] Processing {filename}...")
                
                logger.info(f"[Job {job_id}] Processing: {filename}")
                
                # Extract date from image (uses existing function)
                date_str = extract_date_from_image(img_path)
                
                # Organize by date
                if date_str not in receipts_by_date:
                    receipts_by_date[date_str] = []
                receipts_by_date[date_str].append(img_path)
                
                # Send success update
                send_progress(job_id, f"✅ {filename} → {date_str}")
                
                logger.info(f"[Job {job_id}] Assigned '{date_str}' to {filename}")
                
            except Exception as e:
                logger.error(f"[Job {job_id}] Error processing {img_path}: {e}")
                
                # Add to unknown date
                receipts_by_date.setdefault("Unknown Date", []).append(img_path)
                
                # Send error update
                send_progress(job_id, f"❌ Error processing {filename}: {str(e)}")
        
        # Send document generation status
        send_progress(job_id, f"📄 Creating Word document with {len(receipts_by_date)} date groups...")
        
        # Create document (uses existing function)
        logger.info(f"[Job {job_id}] Creating Word document...")
        doc = create_receipt_document(receipts_by_date)
        
        # Save document
        doc.save(output_doc)
        logger.info(f"[Job {job_id}] Saved Word doc to: {output_doc}")
        
        # Send summary
        summary_lines = [f"  • {date}: {len(files)} receipt(s)" for date, files in sorted(receipts_by_date.items())]
        send_progress(job_id, f"📊 Summary:\n" + "\n".join(summary_lines))
        
        # Cleanup temp files
        cleanup_temp_files()
        
        # Cleanup job temp directory
        try:
            shutil.rmtree(job_tmp_dir, ignore_errors=True)
        except Exception as e:
            logger.warning(f"[Job {job_id}] Could not cleanup temp dir: {e}")
        
        # Send completion with download info
        download_filename = os.path.basename(output_doc)
        send_progress(job_id, f"✅ Word document saved as '{download_filename}'")
        send_progress(job_id, f"🎉 Processing complete! Ready for download.")
        
        logger.info(f"[Job {job_id}] ✅ Processing completed successfully")
        
    except Exception as e:
        logger.error(f"[Job {job_id}] Fatal error: {e}")
        
        # Send error
        send_progress(job_id, f"❌ Processing failed: {str(e)}")
        
        # Cleanup on error
        try:
            shutil.rmtree(job_tmp_dir, ignore_errors=True)
        except:
            pass


# Keep your existing process_receipts function for backward compatibility
def process_receipts(image_paths, output_doc=None):
    """
    Original function - kept for backward compatibility.
    image_paths: list of absolute file paths of images to process
    output_doc: path to save the resulting docx (optional)
    Returns path to saved docx.
    """
    # ensure temp folder exists
    create_temp_folder()

    receipts_by_date = {}
    for img_path in image_paths:
        try:
            logger.info(f"Processing: {os.path.basename(img_path)}")
            date_str = extract_date_from_image(img_path)
            if date_str not in receipts_by_date:
                receipts_by_date[date_str] = []
            receipts_by_date[date_str].append(img_path)
            logger.info(f"Assigned date '{date_str}' to {os.path.basename(img_path)}")
        except Exception as e:
            logger.error(f"Error processing {img_path}: {e}")
            receipts_by_date.setdefault("Unknown Date", []).append(img_path)

    # create document
    doc = create_receipt_document(receipts_by_date)

    # output path
    if not output_doc:
        os.makedirs(os.path.join(os.path.dirname(__file__), '..', 'output'), exist_ok=True)
        output_doc = os.path.join(os.path.dirname(__file__), '..', 'output', CONFIG.get('OUTPUT_DOC', 'Receipts_Sorted.docx'))

    doc.save(output_doc)
    logger.info(f"Saved Word doc to: {output_doc}")

    # cleanup temp files
    cleanup_temp_files()

    return os.path.abspath(output_doc)